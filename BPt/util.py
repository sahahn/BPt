import os
import numpy as np


class BPtInputMixIn():
    pass


def is_array_like(in_val):
    '''Helper function to check if a value is array like or not.

    Parameters
    -----------
    in_val : anything
        The input to check if array like.

    Returns
    ---------
    is_array_like : bool
        True if array like, False if not.

    Examples
    -----------

    .. ipython:: python

        from BPt.util import is_array_like
        import numpy as np

        is_array_like([1, 2, 3])
        is_array_like('some_string')
        is_array_like(5)
        is_array_like(set([1, 2, 3]))
        is_array_like(np.array([1, 2, 3]))

    '''

    if hasattr(in_val, '__len__') and (not isinstance(in_val, str)) and \
       (not isinstance(in_val, dict)) and (not hasattr(in_val, 'fit')) and \
       (not hasattr(in_val, 'transform')):
        return True
    else:
        return False


def conv_to_list(in_val, amt=1):

    if in_val is None:
        return None

    if not is_array_like(in_val) or isinstance(in_val, BPtInputMixIn):
        in_val = [in_val for i in range(amt)]

    return in_val


def save_docx_table(df, filename, decimals=3):
    '''| Helper function for saving a dataframe to a docx file.

    | You must have the library python-docx installed to use
        this function.

    Parameters
    -----------
    df : pandas DataFrame
        The DataFrame in which to save as a docx table.

    filename : str
        A path to a docx file in which to save the table.
        If the file already exists, then the table will be appended to
        the existing file, if it doesn't already exist,
        then it will be created.

    decimals : int, optional
        An optional number of decimal points to
        round any floating point numbers to in the saved
        table.

        ::

            default = 3

    See Also
    ----------
    BPt.Dataset.summary : Create and optionally save a summary of columns.

    Examples
    ---------

    .. ipython:: python

        import BPt as bp
        import pandas as pd

        df = pd.DataFrame(index=['a', 'b', 'c'])
        df.index.name = 'Subject Name'

        df['Column 1'] = [1.00, 2.00, 3.00]
        df['Column 2'] = ['A', 'B', 'C']

        df

        bp.util.save_docx_table(df, 'test.docx')

    Which will save the following table (screenshot taken from table
    opened in LibreOffice):

    .. image:: ../../_static/ex.png
      :alt: Example Saved Table

    '''

    import docx

    # Reset index
    df = df.reset_index()

    if os.path.exists(filename):
        doc = docx.Document(filename)
    else:
        doc = docx.Document()

    t = doc.add_table(df.shape[0]+1, df.shape[1])

    # For each column
    for j in range(df.shape[-1]):

        # Add header
        t.cell(0, j).text = df.columns[j]

        # Check if column is type is float
        col = list(df)[j]
        is_float = 'float' in df[col].dtype.name.lower()

        for i in range(df.shape[0]):

            # Get value
            value = df.values[i, j]

            # Round if float
            if is_float:
                value = np.round(value, decimals=decimals)

            # Set value
            t.cell(i+1, j).text = str(value)

    doc.save(filename)


def substrs(x):
    return {x[i:i+j] for i in range(len(x)) for j in range(len(x) - i + 1)}


def find_substr(data):

    s = substrs(data[0])

    for val in data[1:]:
        s.intersection_update(substrs(val))

    try:
        mx = max(s, key=len)

    except ValueError:
        mx = ''

    return mx


def get_top_substrs(keys):

    found = []
    top = find_substr(keys)

    while len(top) > 1:
        found.append(top)

        keys = [k.replace(top, '') for k in keys]
        top = find_substr(keys)

    return found