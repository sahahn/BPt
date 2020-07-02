import pandas as pd
import numpy as np

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def get_split_df(subject_id, train_subjects, test_subjects):
 
    split_df = pd.DataFrame(columns = [subject_id, 'split'])
    split_df = split_df.set_index(subject_id)

    d = ['Train' for x in range(len(train_subjects))]
    d += ['Test' for x in range(len(test_subjects))]

    i = list(train_subjects) + list(test_subjects)
    split_df['split'] =  pd.Series(data = d, index = i)

    return split_df


def filter_out_nans(dfs):
    
    no_nan_subjects = [dfs[i][~dfs[i].isna().any(axis=1)].index for i in range(len(dfs))]
    
    valid_subjects = no_nan_subjects[0]
    for i in range(1, len(dfs)):
        valid_subjects = np.intersect1d(valid_subjects, no_nan_subjects[i])
        
    dfs = [dfs[i].loc[valid_subjects] for i in range(len(dfs))]
    return dfs


def get_group_by(group_by, dfs, encoders, names, strat_u_name):
    
    if not isinstance(group_by, list):
        group_by = [group_by]
    
    gb_dfs = []
    gb_encoders = []
    gb_names = []

    for g in group_by:

        try:
            ind = names.index(g)
        except ValueError:
            ind = names.index(g + strat_u_name)

        gb_dfs.append(dfs.pop(ind))
        gb_encoders.append(encoders.pop(ind))
        gb_names.append(names.pop(ind))
        
    return gb_dfs, gb_encoders, gb_names


def get_all_df(df):
    
    col_name = list(df)[0]
    all_df = df[[col_name]].copy()
    all_df[col_name] = 'All'
    all_df = all_df.rename({col_name: 'All'}, axis=1)
    all_encoder = {'All': 'All'}

    return [all_df], [all_encoder], ['All']


def proc_group_by(group_by, include_all, dfs, encoders, names, strat_u_name):

    if group_by is not None:
        gb_dfs, gb_encoders, gb_names =\
            get_group_by(group_by, dfs, encoders, names, strat_u_name)
    else:
        gb_dfs, gb_encoders, gb_names = [], [], []

    if group_by is None or include_all:
        all_df, all_encoder, all_name = get_all_df(dfs[0])

        gb_dfs = all_df + gb_dfs
        gb_encoders = all_encoder + gb_encoders
        gb_names = all_name + gb_names
        
    return gb_dfs, gb_encoders, gb_names


def get_subjects_by_group(df, encoder, name):

    if encoder is None:
        raise RuntimeError('Can only group by non-regression/float variables')

    if df.shape[1] == 1:
        subjects_by_group = [(df[name][df[name] == v]).index for v in np.unique(df)]

    else:
        subjects_by_group = [(df[c][df[c] == 1]).index for c in list(df)]
        
    return subjects_by_group


def get_base_title(display_df, name, ind, cat_show_original_name):
    
    if cat_show_original_name:
        title = display_df.loc[ind, 'Original_Name']

        if isinstance(title, int):
            title = name + ' ' + str(title)
        elif isinstance(title, float):
            title = name + ' ' + str(int(title))
        
    elif 'int' in display_df.index.dtype.name:
        title = name + ' ' + str(ind)
    else:
        title = str(ind)
        
    return title


def add_c_entry(table, r, c, text, center=True):
    
    cell = table.rows[r].cells[c]

    if center:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell.paragraphs[0].text = text


def save_table(save_loc, gb_titles, indv_titles, gb_values,
               shape='long', heading='Data Summary', style=None,
               center=True):
    
    doc = Document()

    if heading is not None:
        doc.add_heading(heading)
    
    if shape == 'long':
        
        row_n, col_n = len(indv_titles)+1, len(gb_titles)+1
        table = doc.add_table(rows=row_n, cols=col_n, style=style)
        
        # Add group headers
        for c in range(1, col_n):
            add_c_entry(table, 0, c, gb_titles[c-1])
            
        # Fill in the rest of the rows
        for r in range(1, row_n):
            
            # Add row header
            add_c_entry(table, r, 0, indv_titles[r-1])
            
            # Fill in values
            for c in range(1, col_n):
                add_c_entry(table, r, c, gb_values[c-1][r-1])
    
    # Shape is wide
    else:
        row_n, col_n = len(gb_titles)+1, len(indv_titles)+1
        table = doc.add_table(rows=row_n, cols=col_n, style=style)
        
        # Add indv headers
        for c in range(1, col_n):
            add_c_entry(table, 0, c, indv_titles[c-1])
            
        # Fill in the rest of the rows
        for r in range(1, row_n):
            
            # Add row header
            add_c_entry(table, r, 0, gb_titles[r-1])
            
            # Fill in values
            for c in range(1, col_n):
                add_c_entry(table, r, c, gb_values[r-1][c-1])
    
    # Save
    doc.save(save_loc)